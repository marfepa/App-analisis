"""
Módulo para generación de reportes en Excel y Word.

Este módulo genera:
- Reportes Excel (.xlsx) con tablas, gráficos y formatos condicionales
- Reportes Word (.docx) con resúmenes, listados y recomendaciones
- Incluye resultados de análisis descriptivo y ML
"""

import pandas as pd
import numpy as np
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
import logging

# Excel
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo

# Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    UMBRALES,
    NIVELES_RIESGO,
    COMPETENCIAS,
    REPORTE_CONFIG,
    MENSAJES,
)

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# UTILIDADES
# ============================================================================

def obtener_color_nivel(nivel: str) -> str:
    """
    Obtiene el color hex para un nivel de riesgo.

    Args:
        nivel: Nivel de riesgo ('ALTO', 'MEDIO', 'ALERTA', 'OPTIMO')

    Returns:
        String con color hex (sin '#')
    """
    return NIVELES_RIESGO.get(nivel, {}).get('color', 'FFFFFF')


def obtener_emoji_nivel(nivel: str) -> str:
    """
    Obtiene el emoji para un nivel de riesgo.

    Args:
        nivel: Nivel de riesgo

    Returns:
        Emoji como string
    """
    return NIVELES_RIESGO.get(nivel, {}).get('emoji', '')


# ============================================================================
# GENERACIÓN DE REPORTE EXCEL
# ============================================================================

def crear_hoja_resumen(wb: Workbook,
                       resultados: Dict[str, Any]) -> None:
    """
    Crea hoja de resumen general en Excel.

    Args:
        wb: Workbook de openpyxl
        resultados: Diccionario con resultados del análisis
    """
    ws = wb.create_sheet("Resumen", 0)

    # Estilos
    title_font = Font(name='Calibri', size=16, bold=True, color='FFFFFF')
    title_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font = Font(name='Calibri', size=12, bold=True)
    header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    # Título
    ws['A1'] = 'ANÁLISIS DE RIESGO ACADÉMICO - EDUCACIÓN FÍSICA'
    ws['A1'].font = title_font
    ws['A1'].fill = title_fill
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:F1')
    ws.row_dimensions[1].height = 30

    # Información general
    row = 3
    ws[f'A{row}'] = 'Fecha del análisis:'
    ws[f'B{row}'] = resultados['timestamp']
    ws[f'A{row}'].font = header_font

    row += 1
    ws[f'A{row}'] = 'Total estudiantes:'
    ws[f'B{row}'] = resultados['resumen']['total_estudiantes']
    ws[f'A{row}'].font = header_font

    row += 1
    ws[f'A{row}'] = 'Total cursos:'
    ws[f'B{row}'] = resultados['resumen']['total_cursos']
    ws[f'A{row}'].font = header_font

    # Distribución de riesgo
    row += 2
    ws[f'A{row}'] = 'DISTRIBUCIÓN DE RIESGO'
    ws[f'A{row}'].font = Font(name='Calibri', size=14, bold=True)
    ws[f'A{row}'].fill = header_fill
    ws.merge_cells(f'A{row}:D{row}')

    row += 1
    ws[f'A{row}'] = 'Nivel'
    ws[f'B{row}'] = 'Estudiantes'
    ws[f'C{row}'] = 'Porcentaje'
    ws[f'D{row}'] = 'Estado'
    for col in ['A', 'B', 'C', 'D']:
        ws[f'{col}{row}'].font = header_font
        ws[f'{col}{row}'].fill = header_fill

    # Datos de riesgo
    niveles_data = [
        ('RIESGO ALTO', resultados['resumen']['estudiantes_riesgo_alto'], 'ALTO'),
        ('RIESGO MEDIO', resultados['resumen']['estudiantes_riesgo_medio'], 'MEDIO'),
        ('ALERTA', resultados['resumen']['estudiantes_alerta'], 'ALERTA'),
        ('ÓPTIMO', resultados['resumen']['estudiantes_optimo'], 'OPTIMO'),
    ]

    total = resultados['resumen']['total_estudiantes']

    for nivel_nombre, cantidad, nivel_id in niveles_data:
        row += 1
        ws[f'A{row}'] = nivel_nombre
        ws[f'B{row}'] = cantidad
        ws[f'C{row}'] = f'{cantidad / total * 100:.1f}%' if total > 0 else '0%'
        ws[f'D{row}'] = obtener_emoji_nivel(nivel_id)

        # Aplicar color
        color = obtener_color_nivel(nivel_id)
        fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
        ws[f'A{row}'].fill = fill
        ws[f'B{row}'].fill = fill
        ws[f'C{row}'].fill = fill
        ws[f'D{row}'].fill = fill

    # Ajustar anchos de columna
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 10

    logger.info("  ✓ Hoja 'Resumen' creada")


def crear_hoja_analisis_cursos(wb: Workbook,
                               resultados: Dict[str, Any]) -> None:
    """
    Crea hoja con análisis por curso.

    Args:
        wb: Workbook de openpyxl
        resultados: Diccionario con resultados del análisis
    """
    ws = wb.create_sheet("Análisis por Curso")

    df = resultados['analisis_cursos']

    # Título
    ws['A1'] = 'ANÁLISIS POR CURSO'
    ws['A1'].font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    ws.merge_cells('A1:J1')
    ws.row_dimensions[1].height = 25

    # Escribir datos
    row_offset = 2
    for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), start=row_offset):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)

            # Header style
            if r_idx == row_offset:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
                cell.alignment = Alignment(horizontal='center', vertical='center')

    # Crear tabla
    table_ref = f"A{row_offset}:{chr(65 + len(df.columns) - 1)}{row_offset + len(df)}"
    tab = Table(displayName="TablaCursos", ref=table_ref)
    style = TableStyleInfo(
        name="TableStyleMedium9",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False
    )
    tab.tableStyleInfo = style
    ws.add_table(tab)

    # Ajustar anchos
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = min(max_length + 2, 30)
        ws.column_dimensions[column_letter].width = adjusted_width

    logger.info("  ✓ Hoja 'Análisis por Curso' creada")


def crear_hoja_listado_estudiantes(wb: Workbook,
                                   resultados: Dict[str, Any]) -> None:
    """
    Crea hoja con listado de estudiantes y sus métricas.

    Args:
        wb: Workbook de openpyxl
        resultados: Diccionario con resultados del análisis
    """
    ws = wb.create_sheet("Listado Estudiantes")

    df = resultados['analisis_individual'].copy()

    # Seleccionar columnas relevantes
    columnas_display = [
        'IDEstudiante',
        'nivel_riesgo_final',
        'asist_porcentaje_asistencia',
        'asist_faltas',
        'asist_retrasos',
        'rend_nota_media',
        'rend_n_evaluaciones',
        'asist_tendencia',
        'rend_tendencia',
    ]

    # Agregar columnas ML si existen
    if 'NivelRiesgoML' in df.columns:
        columnas_display.insert(2, 'NivelRiesgoML')
        columnas_display.insert(3, 'Cluster')

    # Filtrar columnas que existen
    columnas_display = [col for col in columnas_display if col in df.columns]
    df_display = df[columnas_display].copy()

    # Renombrar columnas para mejor lectura
    rename_map = {
        'IDEstudiante': 'ID Estudiante',
        'nivel_riesgo_final': 'Nivel Riesgo',
        'NivelRiesgoML': 'Nivel Riesgo ML',
        'Cluster': 'Cluster ML',
        'asist_porcentaje_asistencia': 'Asistencia %',
        'asist_faltas': 'Faltas',
        'asist_retrasos': 'Retrasos',
        'rend_nota_media': 'Nota Media',
        'rend_n_evaluaciones': 'N° Evaluaciones',
        'asist_tendencia': 'Tendencia Asist.',
        'rend_tendencia': 'Tendencia Rend.',
    }
    df_display = df_display.rename(columns=rename_map)

    # Ordenar por nivel de riesgo (más alto primero)
    orden_riesgo = {'ALTO': 0, 'MEDIO': 1, 'ALERTA': 2, 'OPTIMO': 3}
    df_display['_orden'] = df_display['Nivel Riesgo'].map(orden_riesgo)
    df_display = df_display.sort_values('_orden').drop(columns=['_orden'])

    # Título
    ws['A1'] = 'LISTADO DE ESTUDIANTES'
    ws['A1'].font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    ws.merge_cells(f'A1:{chr(65 + len(df_display.columns) - 1)}1')
    ws.row_dimensions[1].height = 25

    # Escribir datos
    row_offset = 2
    for r_idx, row in enumerate(dataframe_to_rows(df_display, index=False, header=True), start=row_offset):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)

            # Header style
            if r_idx == row_offset:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
                cell.alignment = Alignment(horizontal='center', vertical='center')
            else:
                # Formato condicional para nivel de riesgo
                if c_idx == 2:  # Columna de nivel de riesgo
                    nivel = str(value).upper()
                    if nivel in NIVELES_RIESGO:
                        color = obtener_color_nivel(nivel)
                        cell.fill = PatternFill(start_color=color, end_color=color, fill_type='solid')

    # Crear tabla
    table_ref = f"A{row_offset}:{chr(65 + len(df_display.columns) - 1)}{row_offset + len(df_display)}"
    tab = Table(displayName="TablaEstudiantes", ref=table_ref)
    style = TableStyleInfo(
        name="TableStyleMedium9",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False
    )
    tab.tableStyleInfo = style
    ws.add_table(tab)

    # Ajustar anchos
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = min(max_length + 2, 25)
        ws.column_dimensions[column_letter].width = adjusted_width

    logger.info("  ✓ Hoja 'Listado Estudiantes' creada")


def crear_hoja_ml(wb: Workbook,
                 resultados: Dict[str, Any]) -> None:
    """
    Crea hoja con resultados de Machine Learning.

    Args:
        wb: Workbook de openpyxl
        resultados: Diccionario con resultados del análisis
    """
    if resultados.get('ml') is None:
        logger.info("  ⊘ Saltando hoja ML (no hay datos)")
        return

    ws = wb.create_sheet("Machine Learning")

    # Título
    ws['A1'] = 'ANÁLISIS CON MACHINE LEARNING (CLUSTERING)'
    ws['A1'].font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    ws.merge_cells('A1:F1')
    ws.row_dimensions[1].height = 25

    # Explicación
    row = 3
    ws[f'A{row}'] = 'El modelo de Machine Learning agrupa a los estudiantes en clusters basándose en:'
    ws.merge_cells(f'A{row}:F{row}')

    row += 1
    ws[f'A{row}'] = '  • Porcentaje de asistencia y patrón de faltas'
    ws.merge_cells(f'A{row}:F{row}')

    row += 1
    ws[f'A{row}'] = '  • Nota media y tendencias de rendimiento'
    ws.merge_cells(f'A{row}:F{row}')

    row += 1
    ws[f'A{row}'] = '  • Rendimiento por competencias específicas'
    ws.merge_cells(f'A{row}:F{row}')

    # Resultados por cluster
    df_predicciones = resultados['ml']['predicciones']

    row += 2
    ws[f'A{row}'] = 'DISTRIBUCIÓN POR CLUSTER'
    ws[f'A{row}'].font = Font(name='Calibri', size=12, bold=True)
    ws.merge_cells(f'A{row}:D{row}')

    row += 1
    ws[f'A{row}'] = 'Cluster'
    ws[f'B{row}'] = 'Nivel Riesgo'
    ws[f'C{row}'] = 'N° Estudiantes'
    ws[f'D{row}'] = 'Porcentaje'

    for col in ['A', 'B', 'C', 'D']:
        ws[f'{col}{row}'].font = Font(bold=True)
        ws[f'{col}{row}'].fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    # Agrupar por cluster
    cluster_counts = df_predicciones.groupby(['Cluster', 'NivelRiesgoML']).size().reset_index(name='count')
    total_estudiantes = len(df_predicciones)

    for _, cluster_row in cluster_counts.iterrows():
        row += 1
        ws[f'A{row}'] = int(cluster_row['Cluster'])
        ws[f'B{row}'] = cluster_row['NivelRiesgoML']
        ws[f'C{row}'] = int(cluster_row['count'])
        ws[f'D{row}'] = f"{cluster_row['count'] / total_estudiantes * 100:.1f}%"

        # Color según nivel
        color = obtener_color_nivel(cluster_row['NivelRiesgoML'])
        fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
        ws[f'B{row}'].fill = fill

    # Ajustar anchos
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 15

    logger.info("  ✓ Hoja 'Machine Learning' creada")


def generar_reporte_excel(resultados: Dict[str, Any],
                          output_path: str = None) -> str:
    """
    Genera reporte completo en formato Excel.

    Args:
        resultados: Diccionario con resultados del análisis
        output_path: Ruta de salida (None para auto-generar)

    Returns:
        Ruta del archivo generado
    """
    logger.info("\n📊 Generando reporte Excel...")

    # Generar nombre de archivo si no se proporciona
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"outputs/analisis_riesgo_{timestamp}.xlsx"

    # Crear directorio si no existe
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # Crear workbook
    wb = Workbook()

    # Eliminar hoja por defecto
    if 'Sheet' in wb.sheetnames:
        wb.remove(wb['Sheet'])

    # Crear hojas
    crear_hoja_resumen(wb, resultados)
    crear_hoja_analisis_cursos(wb, resultados)
    crear_hoja_listado_estudiantes(wb, resultados)
    crear_hoja_ml(wb, resultados)

    # Guardar
    wb.save(output_path)

    logger.info(f"✅ Reporte Excel generado: {output_path}")

    return output_path


# ============================================================================
# GENERACIÓN DE REPORTE WORD
# ============================================================================

def generar_reporte_word(resultados: Dict[str, Any],
                        output_path: str = None) -> str:
    """
    Genera reporte completo en formato Word.

    Args:
        resultados: Diccionario con resultados del análisis
        output_path: Ruta de salida (None para auto-generar)

    Returns:
        Ruta del archivo generado
    """
    logger.info("\n📄 Generando reporte Word...")

    # Generar nombre de archivo si no se proporciona
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"outputs/informe_riesgo_{timestamp}.docx"

    # Crear directorio si no existe
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # Crear documento
    doc = Document()

    # ========================================================================
    # TÍTULO
    # ========================================================================
    title = doc.add_heading('INFORME DE ANÁLISIS DE RIESGO ACADÉMICO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Educación Física', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fecha
    doc.add_paragraph(f"Fecha del análisis: {resultados['timestamp']}")
    doc.add_paragraph()

    # Advertencia RGPD
    warning = doc.add_paragraph()
    warning.add_run('⚠️ AVISO DE PROTECCIÓN DE DATOS (RGPD)\n').bold = True
    warning.add_run(
        'Este informe utiliza únicamente IDs anonimizados y no contiene '
        'datos personales identificables. Todos los datos han sido procesados '
        'conforme al Reglamento General de Protección de Datos (RGPD).'
    )

    doc.add_page_break()

    # ========================================================================
    # RESUMEN EJECUTIVO
    # ========================================================================
    doc.add_heading('1. RESUMEN EJECUTIVO', 1)

    resumen = resultados['resumen']

    p = doc.add_paragraph()
    p.add_run(f"Total de estudiantes analizados: ").bold = True
    p.add_run(f"{resumen['total_estudiantes']}\n")
    p.add_run(f"Total de cursos: ").bold = True
    p.add_run(f"{resumen['total_cursos']}\n")

    doc.add_heading('Distribución de Riesgo', level=2)

    # Tabla de distribución
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nivel de Riesgo'
    hdr_cells[1].text = 'Estudiantes'
    hdr_cells[2].text = 'Porcentaje'

    # Datos
    niveles_data = [
        (f"{obtener_emoji_nivel('ALTO')} ALTO", resumen['estudiantes_riesgo_alto']),
        (f"{obtener_emoji_nivel('MEDIO')} MEDIO", resumen['estudiantes_riesgo_medio']),
        (f"{obtener_emoji_nivel('ALERTA')} ALERTA", resumen['estudiantes_alerta']),
        (f"{obtener_emoji_nivel('OPTIMO')} ÓPTIMO", resumen['estudiantes_optimo']),
    ]

    total = resumen['total_estudiantes']
    for i, (nivel_text, cantidad) in enumerate(niveles_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = nivel_text
        row_cells[1].text = str(cantidad)
        row_cells[2].text = f"{cantidad / total * 100:.1f}%" if total > 0 else '0%'

    doc.add_page_break()

    # ========================================================================
    # ANÁLISIS POR CURSO
    # ========================================================================
    doc.add_heading('2. ANÁLISIS POR CURSO', 1)

    df_cursos = resultados['analisis_cursos']

    for _, curso in df_cursos.iterrows():
        doc.add_heading(f"Curso: {curso['curso_id']}", level=2)

        # Asistencia
        p = doc.add_paragraph()
        p.add_run('Asistencia:\n').bold = True
        p.add_run(f"  • Media: {curso.get('porcentaje_medio', 0):.1f}%\n")
        p.add_run(f"  • Rango: {curso.get('porcentaje_min', 0):.1f}% - {curso.get('porcentaje_max', 0):.1f}%\n")
        p.add_run(f"  • Estudiantes en riesgo alto: {curso.get('estudiantes_riesgo_alto', 0)}\n")

        # Rendimiento
        p = doc.add_paragraph()
        p.add_run('Rendimiento Académico:\n').bold = True
        p.add_run(f"  • Nota media: {curso.get('rend_nota_media', 0):.2f}\n")
        p.add_run(f"  • Tasa de aprobados: {curso.get('rend_tasa_aprobados', 0):.1f}%\n")
        p.add_run(f"  • Estudiantes en riesgo alto: {curso.get('rend_estudiantes_riesgo_alto', 0)}\n")

    doc.add_page_break()

    # ========================================================================
    # ESTUDIANTES EN RIESGO
    # ========================================================================
    doc.add_heading('3. ESTUDIANTES EN RIESGO', 1)

    df_individual = resultados['analisis_individual']

    # Filtrar por niveles de riesgo
    for nivel in ['ALTO', 'MEDIO']:
        df_nivel = df_individual[df_individual['nivel_riesgo_final'] == nivel]

        if len(df_nivel) > 0:
            doc.add_heading(f"{obtener_emoji_nivel(nivel)} Riesgo {nivel}", level=2)

            doc.add_paragraph(f"Total de estudiantes: {len(df_nivel)}")

            # Lista de estudiantes
            for _, est in df_nivel.iterrows():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"ID: {est['IDEstudiante']} - ").bold = True
                p.add_run(
                    f"Asistencia: {est.get('asist_porcentaje_asistencia', 0):.1f}%, "
                    f"Nota: {est.get('rend_nota_media', 0):.2f}, "
                    f"Tendencia: {est.get('asist_tendencia', 'N/A')}/{est.get('rend_tendencia', 'N/A')}"
                )

    doc.add_page_break()

    # ========================================================================
    # MACHINE LEARNING
    # ========================================================================
    if resultados.get('ml') is not None:
        doc.add_heading('4. ANÁLISIS CON MACHINE LEARNING', 1)

        doc.add_paragraph(
            'Se ha aplicado un modelo de clustering (KMeans) para identificar '
            'automáticamente grupos de estudiantes con características similares. '
            'El modelo considera múltiples factores: asistencia, rendimiento, '
            'tendencias y competencias específicas.'
        )

        df_predicciones = resultados['ml']['predicciones']
        cluster_counts = df_predicciones.groupby(['Cluster', 'NivelRiesgoML']).size().reset_index(name='count')

        doc.add_heading('Clusters Identificados:', level=2)

        for _, cluster_row in cluster_counts.iterrows():
            nivel = cluster_row['NivelRiesgoML']
            p = doc.add_paragraph()
            p.add_run(
                f"{obtener_emoji_nivel(nivel)} Cluster {int(cluster_row['Cluster'])} "
                f"({nivel}): "
            ).bold = True
            p.add_run(f"{int(cluster_row['count'])} estudiantes")

        doc.add_page_break()

    # ========================================================================
    # RECOMENDACIONES
    # ========================================================================
    doc.add_heading('5. RECOMENDACIONES', 1)

    # Recomendaciones basadas en los resultados
    if resumen['estudiantes_riesgo_alto'] > 0:
        doc.add_heading('Acción Inmediata Requerida', level=2)
        p = doc.add_paragraph()
        p.add_run(f"{obtener_emoji_nivel('ALTO')} ").bold = True
        p.add_run(
            f"Hay {resumen['estudiantes_riesgo_alto']} estudiantes en RIESGO ALTO. "
            f"{MENSAJES['recomendaciones']['ALTO_asistencia']}"
        )

    if resumen['estudiantes_riesgo_medio'] > 0:
        doc.add_heading('Seguimiento Estrecho', level=2)
        p = doc.add_paragraph()
        p.add_run(f"{obtener_emoji_nivel('MEDIO')} ").bold = True
        p.add_run(
            f"Hay {resumen['estudiantes_riesgo_medio']} estudiantes en RIESGO MEDIO. "
            f"{MENSAJES['recomendaciones']['MEDIO_asistencia']}"
        )

    # Recomendaciones generales
    doc.add_heading('Recomendaciones Generales', level=2)

    recomendaciones = [
        "Revisar individualmente a cada estudiante en riesgo alto para identificar causas específicas.",
        "Establecer tutorías personalizadas para estudiantes con tendencias negativas.",
        "Comunicar los resultados con las familias de estudiantes en riesgo.",
        "Adaptar metodologías para mejorar la motivación y participación.",
        "Realizar seguimiento quincenal de estudiantes en riesgo medio y alto.",
        "Considerar adaptaciones curriculares cuando sea necesario.",
        "Reforzar competencias específicas donde se detecten debilidades generalizadas.",
    ]

    for rec in recomendaciones:
        doc.add_paragraph(rec, style='List Bullet')

    # Guardar
    doc.save(output_path)

    logger.info(f"✅ Reporte Word generado: {output_path}")

    return output_path


# ============================================================================
# FUNCIÓN PRINCIPAL
# ============================================================================

def generar_reportes(resultados: Dict[str, Any],
                    output_dir: str = 'outputs',
                    formato: str = 'ambos') -> Dict[str, str]:
    """
    Genera reportes en Excel y/o Word.

    Args:
        resultados: Diccionario con resultados del análisis
        output_dir: Directorio de salida
        formato: 'excel', 'word', o 'ambos'

    Returns:
        Diccionario con rutas de archivos generados
    """
    logger.info("\n" + "=" * 70)
    logger.info("📑 GENERANDO REPORTES")
    logger.info("=" * 70)

    # Crear directorio de salida
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    archivos = {}

    # Generar Excel
    if formato in ['excel', 'ambos']:
        excel_path = f"{output_dir}/analisis_riesgo_{timestamp}.xlsx"
        archivos['excel'] = generar_reporte_excel(resultados, excel_path)

    # Generar Word
    if formato in ['word', 'ambos']:
        word_path = f"{output_dir}/informe_riesgo_{timestamp}.docx"
        archivos['word'] = generar_reporte_word(resultados, word_path)

    logger.info("\n" + "=" * 70)
    logger.info("✅ REPORTES GENERADOS EXITOSAMENTE")
    logger.info("=" * 70)

    for tipo, ruta in archivos.items():
        logger.info(f"{tipo.upper()}: {ruta}")

    logger.info("=" * 70 + "\n")

    return archivos


# ============================================================================
# FUNCIÓN DE TEST
# ============================================================================

def test_report_generator():
    """Test básico del generador de reportes."""
    logger.info("🧪 Ejecutando tests de report_generator...")

    # Crear datos de prueba simples
    resultados = {
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'umbrales': UMBRALES,
        'resumen': {
            'total_estudiantes': 10,
            'total_cursos': 2,
            'estudiantes_riesgo_alto': 2,
            'estudiantes_riesgo_medio': 3,
            'estudiantes_alerta': 3,
            'estudiantes_optimo': 2,
        },
        'analisis_cursos': pd.DataFrame({
            'curso_id': ['3ESO-A', '3ESO-B'],
            'estudiantes': [5, 5],
            'porcentaje_medio': [85.0, 90.0],
            'rend_nota_media': [6.5, 7.2],
        }),
        'analisis_individual': pd.DataFrame({
            'IDEstudiante': [f'EST{i:03d}' for i in range(1, 11)],
            'nivel_riesgo_final': ['ALTO', 'ALTO', 'MEDIO', 'MEDIO', 'MEDIO', 'ALERTA', 'ALERTA', 'ALERTA', 'OPTIMO', 'OPTIMO'],
            'asist_porcentaje_asistencia': [70, 72, 80, 82, 85, 88, 90, 92, 95, 98],
            'rend_nota_media': [4.0, 4.5, 5.5, 6.0, 6.5, 7.0, 7.5, 8.0, 8.5, 9.0],
            'asist_tendencia': ['ESTABLE'] * 10,
            'rend_tendencia': ['ESTABLE'] * 10,
        }),
        'ml': None,  # No incluir ML en test simple
    }

    try:
        # Test generación
        archivos = generar_reportes(resultados, output_dir='outputs/test', formato='ambos')

        assert 'excel' in archivos
        assert 'word' in archivos
        assert Path(archivos['excel']).exists()
        assert Path(archivos['word']).exists()

        logger.info("✅ Tests de report_generator PASADOS")
        return True

    except Exception as e:
        logger.error(f"❌ Test falló: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    test_report_generator()
